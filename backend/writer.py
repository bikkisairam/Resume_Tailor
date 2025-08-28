import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

INPUT_FILE = "tailored_resume.json"
OUTPUT_DOCX = "Tailored_Resume.docx"

# -----------------------------
# Styles
# -----------------------------
def configure_styles(doc: Document):
    styles = doc.styles

    # Normal text
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    # Section headings
    if "Section" not in styles:
        section = styles.add_style("Section", WD_STYLE_TYPE.PARAGRAPH)
    else:
        section = styles["Section"]
    section.font.name = "Calibri"
    section.font.size = Pt(12)
    section.font.bold = True
    section.font.color.rgb = RGBColor(0, 0, 0)
    section.paragraph_format.space_before = Pt(0)
    section.paragraph_format.space_after = Pt(0)

    # Subheading (roles, degrees, projects)
    if "SubHeading" not in styles:
        sub = styles.add_style("SubHeading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        sub = styles["SubHeading"]
    sub.font.name = "Calibri"
    sub.font.size = Pt(11)
    sub.font.bold = True
    sub.font.color.rgb = RGBColor(0, 0, 0)
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(0)

def set_margins(doc: Document):
    for s in doc.sections:
        s.top_margin = Inches(0.5)
        s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.5)
        s.right_margin = Inches(0.5)

# -----------------------------
# Helpers
# -----------------------------
def add_line_after(paragraph):
    """Add underline divider below section heading."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    for child in pPr.findall(".//w:pBdr", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
        pPr.remove(child)

    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_section_heading(doc, text):
    p = doc.add_paragraph(text.upper(), style="Section")
    add_line_after(p)

def add_bullets(doc, items):
    for it in items:
        it = it.strip()
        if it:
            doc.add_paragraph(f"• {it}", style="Normal")

def add_tabbed_paragraph(doc, left_text, right_text, style="SubHeading"):
    """Left text + right-aligned date using tab stops based on actual page width."""
    section = doc.sections[0]
    page_width = section.page_width.inches
    left_margin = section.left_margin.inches
    right_margin = section.right_margin.inches
    usable_width = page_width - left_margin - right_margin

    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    tab_stops = p.paragraph_format.tab_stops
    tab_stops.clear_all()
    tab_stops.add_tab_stop(Inches(usable_width), WD_PARAGRAPH_ALIGNMENT.RIGHT)

    run_left = p.add_run(left_text)
    run_left.font.size = Pt(11)
    run_left.bold = True
    run_left.font.name = "Calibri"

    if right_text:
        run_right = p.add_run("\t" + right_text)
        run_right.font.size = Pt(10)
        run_right.font.name = "Calibri"

    return p

# -----------------------------
# Writers
# -----------------------------
def write_header(doc, data):
    details = data.get("Details", {})
    name = details.get("Name", "NAME SURNAME").strip()
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)   # ❌ No underline

    parts = [details.get("Email", "").strip(),
             details.get("Phone", "").strip(),
             details.get("Location", "").strip(),
             details.get("LinkedIn", "").strip(),
             details.get("GitHub", "").strip()]
    parts = [p for p in parts if p]
    if parts:
        line = doc.add_paragraph(" | ".join(parts), style="Normal")
        line.paragraph_format.space_after = Pt(0)

def write_summary(doc, data):
    if data.get("Summary", "").strip():
        add_section_heading(doc, "Summary")
        p = doc.add_paragraph(data["Summary"].strip(), style="Normal")
        p.paragraph_format.space_after = Pt(0)

def write_skills(doc, data):
    skills = data.get("Skills", [])
    if skills:
        add_section_heading(doc, "Skills")
        p = doc.add_paragraph(", ".join(skills), style="Normal")
        p.paragraph_format.space_after = Pt(0)

def write_experience(doc, data):
    exp = data.get("Work Experience", [])
    if not exp: return
    add_section_heading(doc, "Experience")
    for job in exp:
        role = job.get("Role", "").strip()
        company = job.get("Company Name", "").strip()
        date = job.get("Date", "").strip()
        header = f"{company} | {role}" if company or role else ""
        if header:
            add_tabbed_paragraph(doc, header, date, style="SubHeading")
        for b in job.get("Bullet Points", []):
            para = doc.add_paragraph(f"• {b}", style="Normal")
            para.paragraph_format.space_after = Pt(0)

def write_projects(doc, data):
    projects = data.get("Project Experience", [])
    if not projects: return
    add_section_heading(doc, "Projects")
    for p in projects:
        title = p.get("Title", "").strip()
        tech = p.get("Tech Stack", "").strip()
        header = f"{title} | {tech}" if tech else title
        doc.add_paragraph(header, style="SubHeading")
        for b in p.get("Bullet Points", []):
            para = doc.add_paragraph(f"• {b}", style="Normal")
            para.paragraph_format.space_after = Pt(0)

def write_education(doc, data):
    edu = data.get("Education", [])
    if not edu: return
    add_section_heading(doc, "Education")
    for e in edu:
        degree = e.get("Degree", "").strip()
        inst = e.get("Institution", "").strip()
        date = e.get("Date", "").strip()
        gpa = e.get("GPA", "").strip()

        add_tabbed_paragraph(doc, degree, date, style="SubHeading")

        if inst or gpa:
            line2 = inst
            if gpa:
                line2 += f" | GPA: {gpa}"
            para = doc.add_paragraph(line2, style="Normal")
            para.paragraph_format.space_after = Pt(0)

def write_certifications(doc, data):
    certs = data.get("Achievements and Certifications", [])
    if certs:
        add_section_heading(doc, "Achievements & Certifications")
        for c in certs:
            para = doc.add_paragraph(f"• {c}", style="Normal")
            para.paragraph_format.space_after = Pt(0)

# -----------------------------
# Build
# -----------------------------
def build_doc(data: dict, out_path: Path):
    doc = Document()
    set_margins(doc)
    configure_styles(doc)

    write_header(doc, data)
    write_summary(doc, data)
    write_skills(doc, data)
    write_experience(doc, data)
    write_projects(doc, data)
    write_education(doc, data)
    write_certifications(doc, data)

    doc.save(str(out_path))

if __name__ == "__main__":
    data = json.loads(Path(INPUT_FILE).read_text(encoding="utf-8"))
    build_doc(data, Path(OUTPUT_DOCX))
    print(f"✅ DOCX saved to {OUTPUT_DOCX}")
